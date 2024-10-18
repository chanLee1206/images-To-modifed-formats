import os
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import RGBColor
import numpy as np

# Define the path to the directory containing images
src_dir = "src_sample"
# Output Word document path
output_doc = "output.docx"

# Create a new Document
doc = Document()

# Process each image in the specified directory
for filename in os.listdir(src_dir):
    if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):  # You can add more extensions if needed
        # Full path to the image
        img_path = os.path.join(src_dir, filename)
        
        # Open the image using Pillow
        img = Image.open(img_path)

        # Use pytesseract to extract text and word bounding boxes
        d = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

        # Process the data to get text and color
        for i in range(len(d['text'])):
            if int(d['conf'][i]) > 0:  # Only consider positive confidence values
                text = d['text'][i]
                
                # Get bounding box coordinates
                (x, y, w, h) = (d['left'][i], d['top'][i], d['width'][i], d['height'][i])
                
                # Extract the corresponding region of the image to determine color
                region = img.crop((x, y, x + w, y + h))
                
                # Convert to numpy array to analyze the color
                np_region = np.array(region)

                # Ensure the region is not empty and handle small regions
                if np_region.size > 0 and np_region.ndim == 3:  # Ensure it’s an RGB image
                    # Mask for non-white pixels (assuming white background)
                    mask = np.all(np_region[:, :, :3] < 240, axis=2)  # Pixels that are not white
                    if np.any(mask):
                        avg_color = np.mean(np_region[mask], axis=0)[:3]  # Get average RGB values
                        color = RGBColor(int(avg_color[0]), int(avg_color[1]), int(avg_color[2]))
                    else:
                        # Default to black if no valid pixels are found
                        color = RGBColor(0, 0, 0)
                else:
                    # If the region is empty or invalid, set default black color
                    color = RGBColor(0, 0, 0)

                # Add the text with color to the document
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.color.rgb = color  # Set the font color

        # Add a line break after each image's text
        doc.add_paragraph("")  # Adds an empty paragraph for spacing

# Save the document
doc.save(output_doc)

print(f"Converted images in '{src_dir}' to '{output_doc}' successfully!")
